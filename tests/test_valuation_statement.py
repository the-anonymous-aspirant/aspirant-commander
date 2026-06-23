"""Acceptance tests for the Valuation Statement (Värdeutlåtande) tool.

Two layers:
  * unit-level tests of the template populator (no fixtures required)
  * integration-level smoke tests of /extract and /generate; the /extract
    test is skipped when the operator-supplied PDF fixtures aren't
    present (they contain real personnummer + property details so we do
    not commit them to the repo).
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from app.valuation_statement.template import TemplateFields, populate


FIXTURE_ROOT = Path("/tmp/vardeutlatande")
HAS_FIXTURES = (FIXTURE_ROOT / "Datavardering.pdf").exists() and (
    FIXTURE_ROOT / "LGH_utdrag.pdf"
).exists()
HAS_SMAHUS_UC = (FIXTURE_ROOT / "Datavardering_smahus.pdf").exists()
HAS_SMAHUS_PROSE = (FIXTURE_ROOT / "VardeutlatandeHok.pdf").exists()


# ---------- populator unit tests (no PDF fixtures) ----------


def _br_fields() -> TemplateFields:
    return TemplateFields(
        objekt="LGH 1303 HSB Brf Långpannan i Stockholm (7696097448)",
        objekt_short="LGH 1303 HSB Brf Långpannan i Stockholm",
        adress="Hanna Rydhs gata 12",
        kommun="Hägersten",
        upplatelseform="Bostadsrätt",
        datavardering_date="2026-06-02",
        fastighetsutdrag_date=None,
        lagenhetsforteckning_date="2026-06-09",
        bilder_note=None,
        likviditet="normal",
        marknadsvarde_kr="3 050 000",
        intervall_kr="50 000",
        ort="",
        datum="18/6/2026",
        maklare_namn="Jenny Wiklund",
        maklare_titel="Registrerad fastighetsmäklare",
        foretag="Fastighetsbyrån",
        mode="bostadsratt",
    )


def _hok_fields() -> TemplateFields:
    return TemplateFields(
        objekt="Vaggeryd Hok 2:139",
        objekt_short="Vaggeryd Hok 2:139",
        adress="Lillholmsvägen 12",
        kommun="Hok",
        upplatelseform="Friköpt",
        datavardering_date="2026-06-09",
        fastighetsutdrag_date="2026-06-09",
        lagenhetsforteckning_date=None,
        bilder_note="Bilder har inhämtats från kund som underlag.",
        likviditet="normal",
        marknadsvarde_kr="2 200 000",
        intervall_kr="50 000",
        ort="",
        datum="18/6/2026",
        maklare_namn="Jenny Wiklund",
        maklare_titel="Registrerad fastighetsmäklare",
        foretag="Fastighetsbyrån",
        mode="frikopt",
    )


def _paragraph_texts(docx_bytes: bytes) -> list[str]:
    doc = Document(BytesIO(docx_bytes))
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


class TestPopulateBR:
    def test_no_placeholders_remain(self):
        text = "\n".join(_paragraph_texts(populate(_br_fields())))
        assert "[" not in text, "Unfilled placeholder remains in BR output"

    def test_object_row_uses_full_id(self):
        text = _paragraph_texts(populate(_br_fields()))
        assert any(
            p == "Objekt: LGH 1303 HSB Brf Långpannan i Stockholm (7696097448)"
            for p in text
        )

    def test_value_sentence_uses_short_id(self):
        text = "\n".join(_paragraph_texts(populate(_br_fields())))
        assert (
            "Vid värdering av LGH 1303 HSB Brf Långpannan i Stockholm i Hägersten "
            "har ortprismetoden använts."
        ) in text
        # The full id (with orgnr) must NOT appear in the running text.
        assert "Vid värdering av LGH 1303 HSB Brf Långpannan i Stockholm (7696097448)" not in text

    def test_description_drops_fastighetsutdrag_clause(self):
        text = "\n".join(_paragraph_texts(populate(_br_fields())))
        assert (
            "datavärdering per datum 2026-06-02 och "
            "lägenhetsförteckning per datum 2026-06-09 beaktats"
        ) in text
        assert "fastighetsutdrag" not in text

    def test_value_block_uses_supplied_amounts(self):
        text = "\n".join(_paragraph_texts(populate(_br_fields())))
        assert "Marknadsvärdet bedöms till 3 050 000 kr" in text
        assert "intervall om ± 50 000 kr" in text

    def test_empty_ort_yields_datum_only(self):
        # The Ort line should be just the date, not ",18/6/2026" or "Stockholm,18/6/2026".
        paragraphs = [p.strip() for p in _paragraph_texts(populate(_br_fields()))]
        assert "18/6/2026" in paragraphs
        assert not any(p.startswith(",") for p in paragraphs)

    def test_iso_datum_rendered_as_swedish_short_form(self):
        # The Vue client submits the footer date as ISO YYYY-MM-DD from
        # <input type="date">. The populator should render it as DD/M/YYYY.
        fields = _br_fields()
        fields.datum = "2026-06-18"
        paragraphs = [p.strip() for p in _paragraph_texts(populate(fields))]
        assert "18/6/2026" in paragraphs
        assert "2026-06-18" not in "\n".join(paragraphs)

    def test_iso_datum_with_ort_rendered_short_form(self):
        fields = _br_fields()
        fields.ort = "Stockholm"
        fields.datum = "2026-06-18"
        text = "\n".join(_paragraph_texts(populate(fields)))
        assert "Stockholm, 18/6/2026" in text

    def test_marknadsvarde_renders_as_swedish_grouped_when_unspaced(self):
        # Operator types '500000' — the populated output renders '500 000'.
        # Same path is also the regression fix for the '5 000 000 vs 500 000
        # look alike when unspaced' complaint that motivated this work.
        fields = _br_fields()
        fields.marknadsvarde_kr = "5000000"
        fields.intervall_kr = "50000"
        text = "\n".join(_paragraph_texts(populate(fields)))
        assert "Marknadsvärdet bedöms till 5 000 000 kr" in text
        assert "intervall om ± 50 000 kr" in text

    def test_marknadsvarde_already_grouped_passes_through(self):
        # Idempotent: '3 050 000' stays '3 050 000', not '3050 000' or
        # similar from a naive re-format.
        text = "\n".join(_paragraph_texts(populate(_br_fields())))
        assert "Marknadsvärdet bedöms till 3 050 000 kr" in text
        assert "intervall om ± 50 000 kr" in text

    def test_marknadsvarde_paragraph_is_separated_from_liquidity(self):
        # Item 2 of #1026: 'Marknadsvärdet bedöms till …' must be its own
        # paragraph, not run-on after the liquidity sentence.
        paragraphs = _paragraph_texts(populate(_br_fields()))
        liq_idx = next(
            i for i, p in enumerate(paragraphs)
            if "likviditeten bedöms som normal" in p
        )
        mv_idx = next(
            i for i, p in enumerate(paragraphs)
            if p.startswith("Marknadsvärdet bedöms till")
        )
        assert mv_idx == liq_idx + 1, (
            f"Expected Marknadsvärdet paragraph immediately after liquidity "
            f"paragraph, got liq_idx={liq_idx} mv_idx={mv_idx}"
        )

    def test_beskrivning_paragraph_is_unbulleted(self):
        # Item 1 of #1026: the Beskrivning paragraph lands below the four
        # Värderingsobjekt bullets and itself has no bullet marker (numPr).
        from docx.oxml.ns import qn
        doc = Document(BytesIO(populate(_br_fields())))
        beskrivning = next(
            p for p in doc.paragraphs if p.text.startswith("Beskrivning:")
        )
        pPr = beskrivning._p.pPr
        has_numpr = pPr is not None and pPr.find(qn("w:numPr")) is not None
        assert not has_numpr, "Beskrivning paragraph still carries a bullet"

    def test_punctuation_fixes_appear_in_output(self):
        text = "\n".join(_paragraph_texts(populate(_br_fields())))
        # 'objekt Hänsyn' → 'objekt. Hänsyn'
        assert "jämförbara objekt. Hänsyn" in text
        assert "jämförbara objekt Hänsyn" not in text
        # 'K/T värde' → 'K/T-värde'
        assert "K/T-värde" in text
        assert "K/T värde" not in text
        # 'Underlag för kreditprövning' — no trailing space.
        assert "Underlag för kreditprövning\n" in text + "\n"
        assert "Underlag för kreditprövning " not in text

    def test_populated_output_has_no_highlights(self):
        # The mall.docx ships with placeholders yellow-highlighted as an
        # authoring aid. The footer date is the visible failure mode (its
        # first run carries the highlight, which would otherwise survive
        # substitution), but every populated field must render unhighlighted.
        doc = Document(BytesIO(populate(_br_fields())))
        offenders = [
            (pi, ri, run.text, run.font.highlight_color)
            for pi, para in enumerate(doc.paragraphs)
            for ri, run in enumerate(para.runs)
            if run.font.highlight_color is not None and run.text
        ]
        assert offenders == [], f"populated runs still carry highlights: {offenders}"


class TestPopulateHok:
    def test_object_row_uses_kommun_fastighet(self):
        paragraphs = _paragraph_texts(populate(_hok_fields()))
        assert "Objekt: Vaggeryd Hok 2:139" in paragraphs

    def test_description_drops_lagenhetsforteckning_clause(self):
        text = "\n".join(_paragraph_texts(populate(_hok_fields())))
        assert (
            "datavärdering per datum 2026-06-09 och "
            "fastighetsutdrag per datum 2026-06-09 beaktats"
        ) in text
        assert "lägenhetsförteckning" not in text

    def test_bilder_note_inserted(self):
        text = "\n".join(_paragraph_texts(populate(_hok_fields())))
        assert "Bilder har inhämtats från kund som underlag." in text

    def test_no_metacomments_remain(self):
        text = "\n".join(_paragraph_texts(populate(_hok_fields())))
        assert "Enbart bostadsrätt" not in text
        assert "Fyll på ifall bilder" not in text


# ---------- /generate API smoke test (no PDF fixtures needed) ----------


def test_generate_endpoint_returns_docx(client):
    body = {
        "objekt": "LGH 1303 HSB Brf Långpannan i Stockholm (7696097448)",
        "objekt_short": "LGH 1303 HSB Brf Långpannan i Stockholm",
        "adress": "Hanna Rydhs gata 12",
        "kommun": "Hägersten",
        "upplatelseform": "Bostadsrätt",
        "mode": "bostadsratt",
        "datavardering_date": "2026-06-02",
        "lagenhetsforteckning_date": "2026-06-09",
        "likviditet": "normal",
        "marknadsvarde_kr": "3 050 000",
        "intervall_kr": "50 000",
        "ort": "",
        "datum": "18/6/2026",
        "maklare_namn": "Jenny Wiklund",
        "maklare_titel": "Registrerad fastighetsmäklare",
        "foretag": "Fastighetsbyrån",
    }
    r = client.post("/valuation-statement/generate", json=body)
    assert r.status_code == 200
    assert (
        "officedocument.wordprocessingml.document" in r.headers["content-type"]
    )
    assert r.headers["content-disposition"].startswith("attachment;")
    assert r.content[:4] == b"PK\x03\x04"  # docx is a ZIP

    text = "\n".join(_paragraph_texts(r.content))
    assert "Objekt: LGH 1303 HSB Brf Långpannan i Stockholm (7696097448)" in text


# ---------- fastighetsutdrag dispatch (no PDF fixture required) ----------


def test_extract_document_dispatches_fastighetsutdrag_without_importerror():
    """Guard the dispatch table against missing parser modules.

    The classifier recognises FASTIGHETSUTDRAG, so the route hands one
    to `extract_document()`. If the per-type parser module is missing
    the call raises ImportError and the whole `/extract` request 500s.
    The stub parser returns an empty result — the operator still types
    `fastighetsutdrag_date` during review — but the dispatch must not
    raise.
    """
    from app.valuation_statement.classifier import DocumentType
    from app.valuation_statement.extraction import extract_document

    result = extract_document(b"", DocumentType.FASTIGHETSUTDRAG, "utdrag.pdf")

    assert result.document_type == DocumentType.FASTIGHETSUTDRAG
    assert result.filename == "utdrag.pdf"
    assert result.fields == []


# ---------- /extract integration (skip when fixtures missing) ----------


@pytest.mark.skipif(not HAS_FIXTURES, reason="Operator-supplied PDF samples not present")
def test_extract_endpoint_parses_both_types(client):
    with (FIXTURE_ROOT / "Datavardering.pdf").open("rb") as a, (
        FIXTURE_ROOT / "LGH_utdrag.pdf"
    ).open("rb") as b:
        r = client.post(
            "/valuation-statement/extract",
            files=[
                ("files", ("Datavärdering.pdf", a.read(), "application/pdf")),
                ("files", ("LGH utdrag.pdf", b.read(), "application/pdf")),
            ],
        )
    assert r.status_code == 200
    data = r.json()
    types = {d["document_type"] for d in data["documents"]}
    assert types == {"datavardering_br", "lgh_utdrag"}

    fields_by_key = {
        d["document_type"]: {f["key"]: f["value"] for f in d["fields"]}
        for d in data["documents"]
    }
    assert fields_by_key["datavardering_br"]["forening_namn"] == (
        "HSB Brf Långpannan i Stockholm"
    )
    assert fields_by_key["datavardering_br"]["marknadsvarde_suggested"] == "2 350 000"
    assert fields_by_key["lgh_utdrag"]["lgh_skatteverket"] == "1303"
    assert fields_by_key["lgh_utdrag"]["postort"] == "Hägersten"


# ---------- comparable-sales row parser (unit, no PDF) ----------


class TestComparableRowParser:
    """The UC Bostad PDF has no column borders, so we parse the flat-text
    rows by anchoring on the trailing YYYY-MM and walking right-to-left.
    These tests pin both shapes that show up in the live sample."""

    def _parse(self, line):
        from app.valuation_statement.parsers.bostadsratt import _parse_comparable_row
        return _parse_comparable_row(line)

    def test_eight_column_row_with_balkong(self):
        row = self._parse("HSBBrfLångpannaniStockholm 72 Ja 6307 302610 2920000 40556 2026-04")
        assert row == {
            "forening": "HSB Brf Långpannan i Stockholm",
            "area_m2": "72",
            "balkong": "Ja",
            "avgift_kr_manad": "6307",
            "arsavgift_kr": "302610",
            "pris_kr": "2920000",
            "pris_per_m2": "40556",
            "salj_datum": "2026-04",
            "raw": "HSBBrfLångpannaniStockholm 72 Ja 6307 302610 2920000 40556 2026-04",
        }

    def test_seven_column_row_no_balkong(self):
        row = self._parse("HSBBrfLångpannaniStockholm 62,5 5002 262682 2400000 38400 2025-12")
        assert row["balkong"] is None
        assert row["area_m2"] == "62,5"
        assert row["avgift_kr_manad"] == "5002"
        assert row["pris_kr"] == "2400000"
        assert row["salj_datum"] == "2025-12"

    def test_row_with_balkong_nej(self):
        row = self._parse("HSBBrfLångpannaniStockholm 70 Nej 5500 250000 2500000 35714 2026-02")
        assert row["balkong"] == "Nej"
        assert row["area_m2"] == "70"

    def test_line_without_date_is_skipped(self):
        assert self._parse("Adress Boyta Avgift Pris Datum") is None

    def test_too_few_columns_falls_back_to_raw(self):
        # Malformed row — at least the raw line + the trailing date survive
        # so the operator still sees something rather than the row vanishing.
        row = self._parse("Onlyforening 2026-04")
        assert row is not None
        assert row["salj_datum"] == "2026-04"
        assert row["raw"] == "Onlyforening 2026-04"
        assert row.get("forening") is None


# ---------- Småhus parser (skip when sample PDF missing) ----------


@pytest.mark.skipif(not HAS_SMAHUS_UC, reason="Operator UC Småhus sample missing")
def test_extract_endpoint_parses_uc_smahus_tabular(client):
    """The UC Bostad Småhus data-feed report populates the six-column
    Objektsinformation band slots plus the two-column Värde block. The
    regression #1059 referenced is the silent-empty case where the
    classifier returned UNKNOWN and `fields[]` was empty — this asserts
    the parser is wired and the column-anchor walk picks up every slot
    the UC sample exposes.
    """
    with (FIXTURE_ROOT / "Datavardering_smahus.pdf").open("rb") as a:
        r = client.post(
            "/valuation-statement/extract",
            files=[("files", ("Datavardering_smahus.pdf", a.read(), "application/pdf"))],
        )
    assert r.status_code == 200
    docs = r.json()["documents"]
    assert len(docs) == 1
    doc = docs[0]
    assert doc["document_type"] == "datavardering_smahus"

    fields = {f["key"]: f for f in doc["fields"]}
    assert fields["fastighetsbeteckning"]["value"] == "Bengtsfors Närsidan 1:21"
    assert fields["fastighetsbeteckning"]["confidence"] == "confident"
    assert fields["byggnadstyp"]["value"] == "Friliggande"
    assert fields["taxeringsvarde"]["value"] == "765 000 (Tax24)"
    assert fields["tomtyta_m2"]["value"] == "2 089"
    assert fields["byggnadsar"]["value"] == "1980"
    assert fields["vardear"]["value"] == "1980"
    assert fields["marknadsvarde_suggested"]["value"] == "1 075 000"
    assert fields["osakerhet_uppat"]["value"] == "350 000"
    assert fields["osakerhet_nedat"]["value"] == "280 000"
    assert fields["document_date"]["value"] == "2026-06-18"


@pytest.mark.skipif(not HAS_SMAHUS_PROSE, reason="Operator Fastighetsbyrån prose Småhus sample missing")
def test_extract_endpoint_parses_prose_smahus_appraisal(client):
    """The Fastighetsbyrån Värdeutlåtande prose layout (sample
    VardeutlatandeHok.pdf) classifies into the same DocumentType as the
    UC tabular report and feeds through the same parser branch — per
    the operator correction on epic #1060 (content-only, one parser
    branch per content type regardless of issuer). The prose layout
    exposes a narrower slot set; the slots it does carry must surface
    with confident labels.
    """
    with (FIXTURE_ROOT / "VardeutlatandeHok.pdf").open("rb") as a:
        r = client.post(
            "/valuation-statement/extract",
            files=[("files", ("VardeutlatandeHok.pdf", a.read(), "application/pdf"))],
        )
    assert r.status_code == 200
    docs = r.json()["documents"]
    assert len(docs) == 1
    doc = docs[0]
    assert doc["document_type"] == "datavardering_smahus"

    fields = {f["key"]: f for f in doc["fields"]}
    assert fields["fastighetsbeteckning"]["value"] == "Vaggeryd Hok 2:139"
    assert fields["adress"]["value"] == "Lillholmsvägen 12"
    assert fields["kommun"]["value"] == "Hok"
    assert fields["upplatelseform"]["value"] == "Friköpt"
    assert fields["marknadsvarde_suggested"]["value"] == "2 200 000"
    # Prose appraisals report a symmetric ±X kr interval — uppåt and
    # nedåt must both carry the same X.
    assert fields["osakerhet_uppat"]["value"] == "50 000"
    assert fields["osakerhet_nedat"]["value"] == "50 000"
    assert fields["document_date"]["value"] == "2026-06-18"
    for slot in (
        "fastighetsbeteckning",
        "adress",
        "kommun",
        "upplatelseform",
        "marknadsvarde_suggested",
        "document_date",
    ):
        assert fields[slot]["confidence"] == "confident", f"{slot} should be confident"


# ---------- prose-extraction regex (unit, no PDF) ----------


class TestProseSmahusRegex:
    """Unit-level guards for the prose-layout regexes so they don't
    silently regress against pdfplumber output variations."""

    def _parse(self, text):
        # Build a minimal ParseContext from page-1 text only — the prose
        # strategies don't need word coordinates.
        from app.valuation_statement.parsers._context import ParseContext
        from app.valuation_statement.parsers._strategy import run_slots
        from app.valuation_statement.parsers.smahus import _SLOTS

        ctx = ParseContext(
            page1_text=text,
            page1_words=(),
            page_texts=(text,),
        )
        return {f.key: f.value for f in run_slots(_SLOTS, ctx, "x.pdf")}

    def test_bullets_with_visible_bullet_glyph(self):
        text = (
            "Värderingsobjekt\n"
            "● Objekt: Vaggeryd Hok 2:139\n"
            "● Adress: Lillholmsvägen 12\n"
            "● Kommun: Hok\n"
            "● Upplåtelseform: Friköpt\n"
            "Marknadsvärdet bedöms till 2 200 000 kr, med ett intervall om +/- 50 000 kr.\n"
            "18/6/2026\n"
        )
        fields = self._parse(text)
        assert fields["fastighetsbeteckning"] == "Vaggeryd Hok 2:139"
        assert fields["adress"] == "Lillholmsvägen 12"
        assert fields["kommun"] == "Hok"
        assert fields["upplatelseform"] == "Friköpt"
        assert fields["marknadsvarde_suggested"] == "2 200 000"
        assert fields["osakerhet_uppat"] == "50 000"
        assert fields["osakerhet_nedat"] == "50 000"
        assert fields["document_date"] == "2026-06-18"

    def test_bullets_collapsed_without_glyph(self):
        # The BR prose sample uses ± instead of +/- and the bullet
        # glyph may be stripped by pdfplumber depending on font.
        text = (
            "Värderingsobjekt\n"
            "Objekt: Vaggeryd Hok 2:139\n"
            "Adress: Lillholmsvägen 12\n"
            "Marknadsvärdet bedöms till 2 200 000 kr, med ett intervall om ± 50 000 kr.\n"
            "1/12/2026\n"
        )
        fields = self._parse(text)
        assert fields["fastighetsbeteckning"] == "Vaggeryd Hok 2:139"
        assert fields["adress"] == "Lillholmsvägen 12"
        assert fields["marknadsvarde_suggested"] == "2 200 000"
        assert fields["osakerhet_uppat"] == "50 000"
        assert fields["document_date"] == "2026-12-01"

    def test_missing_value_sentence_falls_back_to_not_found(self):
        text = "Värderingsobjekt\n● Objekt: Foo 1:1\n"
        fields = self._parse(text)
        assert fields["fastighetsbeteckning"] == "Foo 1:1"
        assert fields["marknadsvarde_suggested"] is None
        assert fields["osakerhet_uppat"] is None


@pytest.mark.skipif(not HAS_FIXTURES, reason="Operator-supplied PDF samples not present")
def test_extract_endpoint_returns_structured_comparable_sales(client):
    """The /extract response carries the per-row column dict, not just `raw`."""
    with (FIXTURE_ROOT / "Datavardering.pdf").open("rb") as a:
        r = client.post(
            "/valuation-statement/extract",
            files=[("files", ("Datavärdering.pdf", a.read(), "application/pdf"))],
        )
    assert r.status_code == 200
    docs = r.json()["documents"]
    dv = next(d for d in docs if d["document_type"] == "datavardering_br")
    assert dv["comparable_sales"], "Expected at least one parsed comparable row"
    first = dv["comparable_sales"][0]
    # Structured columns are present on every row (some may be None).
    for col in ("forening", "area_m2", "pris_kr", "pris_per_m2", "salj_datum", "raw"):
        assert col in first


# ---------- operator defaults ----------


def test_operator_defaults_first_load_returns_example_values(client, tmp_path, monkeypatch):
    """No file yet → GET returns the ground-truth identity so the review
    step is pre-filled on first deploy instead of blank."""
    monkeypatch.setenv(
        "VALUATION_OPERATOR_DEFAULTS_PATH", str(tmp_path / "valuation_defaults.json")
    )
    r = client.get("/valuation-statement/operator-defaults")
    assert r.status_code == 200
    defaults = r.json()
    assert defaults["maklare_namn"] == "Jenny Wiklund"
    assert defaults["maklare_titel"] == "Registrerad fastighetsmäklare"
    assert defaults["foretag"] == "Fastighetsbyrån"
    assert defaults["ort"] == "Nynäshamn"
    assert defaults["likviditet"] == "normal"


def test_operator_defaults_saved_values_override_examples(client, tmp_path, monkeypatch):
    """After PUT, GET reflects the saved values — including a blank ort,
    which is a legitimate operator choice for the 'date-only' footer."""
    path = tmp_path / "valuation_defaults.json"
    monkeypatch.setenv("VALUATION_OPERATOR_DEFAULTS_PATH", str(path))

    saved = {
        "ort": "",
        "maklare_namn": "Anna Andersson",
        "maklare_titel": "Mäklare",
        "foretag": "Annas Mäkleri",
        "likviditet": "god",
    }
    put = client.put("/valuation-statement/operator-defaults", json=saved)
    assert put.status_code == 200
    assert path.exists()

    r = client.get("/valuation-statement/operator-defaults")
    assert r.status_code == 200
    defaults = r.json()
    assert defaults["maklare_namn"] == "Anna Andersson"
    assert defaults["ort"] == ""
    assert defaults["likviditet"] == "god"


# ---------- PDF export ----------


_GENERATE_BODY = {
    "objekt": "LGH 1303 HSB Brf Långpannan i Stockholm (7696097448)",
    "objekt_short": "LGH 1303 HSB Brf Långpannan i Stockholm",
    "adress": "Hanna Rydhs gata 12",
    "kommun": "Hägersten",
    "upplatelseform": "Bostadsrätt",
    "mode": "bostadsratt",
    "datavardering_date": "2026-06-02",
    "lagenhetsforteckning_date": "2026-06-09",
    "likviditet": "normal",
    "marknadsvarde_kr": "3 050 000",
    "intervall_kr": "50 000",
    "ort": "",
    "datum": "18/6/2026",
    "maklare_namn": "Jenny Wiklund",
    "maklare_titel": "Registrerad fastighetsmäklare",
    "foretag": "Fastighetsbyrån",
}


def test_generate_pdf_calls_libreoffice(client, monkeypatch):
    from app.valuation_statement import routes as v_routes

    captured: dict = {}

    def fake_docx_to_pdf(docx_bytes: bytes) -> bytes:
        captured["docx_len"] = len(docx_bytes)
        return b"%PDF-1.7\n<<fake pdf body>>\n%%EOF"

    monkeypatch.setattr(v_routes, "docx_to_pdf", fake_docx_to_pdf)

    r = client.post("/valuation-statement/generate?format=pdf", json=_GENERATE_BODY)
    assert r.status_code == 200
    assert r.headers["content-type"] == "application/pdf"
    assert r.content.startswith(b"%PDF")
    assert r.headers["content-disposition"].endswith('.pdf"')
    assert captured["docx_len"] > 0


def test_generate_pdf_503_when_libreoffice_missing(client, monkeypatch):
    from app.valuation_statement import routes as v_routes
    from app.valuation_statement.pdf_export import LibreOfficeUnavailable

    def raise_unavailable(_docx_bytes):
        raise LibreOfficeUnavailable("soffice not on PATH")

    monkeypatch.setattr(v_routes, "docx_to_pdf", raise_unavailable)
    r = client.post("/valuation-statement/generate?format=pdf", json=_GENERATE_BODY)
    assert r.status_code == 503
    assert "soffice" in r.json()["detail"].lower()


def test_generate_default_format_is_docx(client):
    """No ?format=… → default 'docx' (existing tests rely on this)."""
    r = client.post("/valuation-statement/generate", json=_GENERATE_BODY)
    assert r.status_code == 200
    assert "officedocument.wordprocessingml.document" in r.headers["content-type"]


def test_generate_rejects_unknown_format(client):
    r = client.post("/valuation-statement/generate?format=ps", json=_GENERATE_BODY)
    assert r.status_code == 422
